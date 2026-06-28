from __future__ import annotations

from io import BytesIO
import textwrap
from datetime import timedelta
from typing import Any
from uuid import UUID

from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.analysis import Analysis
from app.models.base import utc_now
from app.models.export import Export
from app.models.project import Project
from app.utils.file_validation import build_export_object_key, safe_filename

EXPORT_RETENTION_DAYS = 7
MARKDOWN_MIME_TYPE = "text/markdown; charset=utf-8"
PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXPORT_MIME_TYPES = {
    "markdown": MARKDOWN_MIME_TYPE,
    "pdf": PDF_MIME_TYPE,
    "docx": DOCX_MIME_TYPE,
}
EXPORT_EXTENSIONS = {
    "markdown": "md",
    "pdf": "pdf",
    "docx": "docx",
}
REPORT_TOC = [
    "Ringkasan Proyek",
    "Readiness Score",
    "Penjelasan Skor",
    "Rekomendasi Langkah Berikutnya",
    "Checklist Revisi Skripsi",
    "Checklist Temuan & Perbaikan",
    "Konsistensi Slide vs Laporan",
    "Klaim Bermasalah",
    "Pertanyaan Penguji",
    "Script Presentasi",
    "Disclaimer",
]


def list_exports_for_project(
    db: Session,
    project_id: UUID,
    sort: str = "latest",
    export_type: str | None = None,
) -> list[Export]:
    order_by = Export.created_at.asc() if sort == "oldest" else Export.created_at.desc()
    statement = select(Export).where(Export.project_id == project_id)
    if export_type:
        statement = statement.where(Export.export_type == export_type)
    statement = statement.order_by(order_by)
    return list(db.execute(statement).scalars().all())


def get_export_for_project(db: Session, export_id: UUID, project_id: UUID) -> Export:
    statement = select(Export).where(
        Export.id == export_id,
        Export.project_id == project_id,
    )
    export = db.execute(statement).scalar_one_or_none()
    if export is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Export tidak ditemukan.",
        )
    return export


def create_export_record(
    db: Session,
    *,
    project: Project,
    analysis: Analysis,
    user_id: UUID,
    export_type: str,
) -> tuple[Export, bytes]:
    if analysis.status != "success" or not analysis.result_json:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Export hanya dapat dibuat setelah analisis selesai.",
        )
    if export_type not in EXPORT_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Format export tidak didukung.",
        )

    created_at = utc_now()
    markdown = render_markdown_report(project=project, result=dict(analysis.result_json))
    file_bytes = render_export_bytes(markdown, export_type)
    timestamp = created_at.strftime("%Y%m%d-%H%M%S")
    extension = EXPORT_EXTENSIONS[export_type]
    file_name = (
        f"{safe_filename(project.title).lower()}-readiness-report-"
        f"{timestamp}.{extension}"
    )
    export = Export(
        project_id=project.id,
        export_type=export_type,
        file_name=file_name,
        file_mime_type=EXPORT_MIME_TYPES[export_type],
        file_size=len(file_bytes),
        r2_object_key="pending",
        created_at=created_at,
        updated_at=created_at,
        expires_at=created_at + timedelta(days=EXPORT_RETENTION_DAYS),
    )
    db.add(export)
    db.flush()
    export.r2_object_key = build_export_object_key(
        user_id=user_id,
        project_id=project.id,
        export_id=export.id,
        file_name=file_name,
    )
    return export, file_bytes


def render_export_bytes(markdown: str, export_type: str) -> bytes:
    if export_type == "markdown":
        return markdown.encode("utf-8")
    if export_type == "pdf":
        return render_pdf_report(markdown)
    if export_type == "docx":
        return render_docx_report(markdown)
    raise ValueError(f"Unsupported export type: {export_type}")


def render_markdown_report(*, project: Project, result: dict[str, Any]) -> str:
    overview = dict(result.get("overview") or {})
    lines = [
        f"# Laporan Kesiapan Sidang - {project.title}",
        "",
        "## Daftar Isi",
        "",
        *[f"- [{title}](#{_markdown_anchor(title)})" for title in REPORT_TOC],
        "",
        "## Ringkasan Proyek",
        "",
        f"- Judul skripsi: {project.thesis_title}",
        f"- Mahasiswa: {project.student_name}",
        f"- Program studi: {project.major}",
        f"- Universitas/kampus: {project.university}",
        f"- Target durasi presentasi: {project.target_presentation_minutes} menit",
        "",
        "## Readiness Score",
        "",
        f"- Skor: {overview.get('readiness_score', 0)}/100",
        f"- Kategori: {overview.get('readiness_category', 'Belum tersedia')}",
        f"- Isu kritis: {overview.get('critical_issues_count', 0)}",
        f"- Isu penting: {overview.get('important_issues_count', 0)}",
        f"- Isu minor: {overview.get('minor_issues_count', 0)}",
        f"- Total slide dicek: {overview.get('total_slides_checked', 0)}",
        f"- Total pertanyaan penguji: {overview.get('total_examiner_questions', 0)}",
        "",
    ]

    lines.extend(_section_list("Penjelasan Skor", overview.get("score_explanation")))
    lines.extend(
        _section_list("Rekomendasi Langkah Berikutnya", overview.get("recommended_next_actions"))
    )
    lines.extend(_revision_section("Checklist Revisi Skripsi", result.get("official_revision_items")))
    lines.extend(_revision_section("Checklist Temuan & Perbaikan", result.get("revision_items")))
    lines.extend(_slide_consistency_section(result.get("slide_checks")))
    lines.extend(_problematic_claims_section(result.get("problematic_claims")))
    lines.extend(_defense_questions_section(result.get("defense_questions")))
    lines.extend(_presentation_script_section(result.get("presentation_scripts")))

    disclaimer = (
        overview.get("disclaimer")
        or result.get("raw_disclaimer")
        or "Hasil ini adalah saran berbasis dokumen yang diunggah dan perlu ditinjau ulang bersama dosen pembimbing."
    )
    lines.extend(
        [
            "## Disclaimer",
            "",
            _text(disclaimer),
            "",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def render_pdf_report(markdown: str) -> bytes:
    import fitz

    document = fitz.open()
    page = document.new_page()
    margin = 48
    max_width = 86
    y = margin
    line_height = 13
    page_height = page.rect.height
    toc_links = []
    heading_targets = {}

    def ensure_space(required: float) -> None:
        nonlocal page, y
        if y + required > page_height - margin:
            page = document.new_page()
            y = margin

    for raw_line in markdown.splitlines():
        font_size = 11
        indent = 0
        toc_title = None
        text = raw_line.strip()
        if raw_line.startswith("# "):
            font_size = 18
            text = text.removeprefix("# ").strip()
            ensure_space(34)
            y += 6
        elif raw_line.startswith("## "):
            font_size = 15
            text = text.removeprefix("## ").strip()
            ensure_space(28)
            y += 8
            if text in REPORT_TOC:
                heading_targets[text] = (len(document) - 1, y)
        elif raw_line.startswith("### "):
            font_size = 13
            text = text.removeprefix("### ").strip()
            ensure_space(24)
            y += 4
        elif raw_line.startswith("- "):
            indent = 14
            list_text = text.removeprefix("- ").strip()
            parsed_link = _parse_markdown_link(list_text)
            if parsed_link and parsed_link[0] in REPORT_TOC:
                toc_title = parsed_link[0]
            text = f"- {_strip_markdown_link(list_text)}"
        elif text.startswith("**") and text.endswith(":**"):
            text = text.strip("*")
            font_size = 12

        wrapped_lines = textwrap.wrap(text, width=max_width - int(indent / 2)) if text else [""]
        for line in wrapped_lines:
            ensure_space(line_height + max(0, font_size - 11) + 2)
            if line:
                if toc_title and line.startswith("- "):
                    toc_links.append(
                        (
                            len(document) - 1,
                            fitz.Rect(
                                margin + indent,
                                y - font_size,
                                page.rect.width - margin,
                                y + 4,
                            ),
                            toc_title,
                        )
                    )
                page.insert_text(
                    (margin + indent, y),
                    line,
                    fontsize=font_size,
                    fontname="helv",
                )
            y += line_height + max(0, font_size - 11)

    for page_index, rect, title in toc_links:
        target = heading_targets.get(title)
        if not target:
            continue
        target_page_index, target_y = target
        document[page_index].insert_link(
            {
                "kind": fitz.LINK_GOTO,
                "from": rect,
                "page": target_page_index,
                "to": fitz.Point(margin, max(target_y - 18, margin)),
            }
        )

    for index, pdf_page in enumerate(document, start=1):
        footer = f"SidangReady AI - {index}"
        pdf_page.insert_text(
            (margin, pdf_page.rect.height - 24),
            footer,
            fontsize=8,
            color=(0.35, 0.39, 0.48),
        )

    pdf_bytes = document.tobytes()
    document.close()
    return pdf_bytes


def render_docx_report(markdown: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    document = DocxDocument()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    bookmark_id = 0

    def add_bookmark(paragraph: Any, name: str) -> None:
        nonlocal bookmark_id
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(bookmark_id))
        bookmark_start.set(qn("w:name"), name)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(bookmark_id))
        insert_at = 1 if paragraph._p.pPr is not None else 0
        paragraph._p.insert(insert_at, bookmark_start)
        paragraph._p.append(bookmark_end)
        bookmark_id += 1

    def add_internal_hyperlink(paragraph: Any, text: str, bookmark_name: str) -> None:
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)
        hyperlink.set(qn("w:history"), "1")

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        style = OxmlElement("w:rStyle")
        style.set(qn("w:val"), "Hyperlink")
        run_properties.append(style)
        run.append(run_properties)

        run_text = OxmlElement("w:t")
        run_text.text = text
        run.append(run_text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
        elif line.startswith("# "):
            document.add_heading(line.removeprefix("# ").strip(), level=1)
        elif line.startswith("## "):
            title = line.removeprefix("## ").strip()
            paragraph = document.add_heading(title, level=2)
            if title in REPORT_TOC:
                add_bookmark(paragraph, _bookmark_name(title))
        elif line.startswith("### "):
            document.add_heading(line.removeprefix("### ").strip(), level=3)
        elif line.startswith("- "):
            link = _parse_markdown_link(line.removeprefix("- ").strip())
            paragraph = document.add_paragraph(style="List Bullet")
            if link and link[0] in REPORT_TOC:
                add_internal_hyperlink(paragraph, link[0], _bookmark_name(link[0]))
            else:
                paragraph.add_run(_strip_markdown_link(line.removeprefix("- ").strip()))
        elif line.startswith("**") and line.endswith(":**"):
            paragraph = document.add_paragraph()
            paragraph.add_run(line.strip("*")).bold = True
        else:
            document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _section_list(title: str, items: object) -> list[str]:
    values = _as_list(items)
    lines = [f"## {title}", ""]
    if not values:
        return lines + ["Belum ada data.", ""]
    return lines + [f"- {_text(item)}" for item in values] + [""]


def _revision_section(title: str, items: object) -> list[str]:
    values = [item for item in _as_list(items) if isinstance(item, dict)]
    lines = [f"## {title}", ""]
    if not values:
        return lines + ["Belum ada item.", ""]
    for index, item in enumerate(values, start=1):
        related = _related_text(item)
        lines.extend(
            [
                f"### {index}. {_text(item.get('title'))}",
                "",
                f"- Prioritas: {_text(item.get('priority'))}",
                f"- Status: {_text(item.get('status'))}",
                f"- Terkait: {related}",
                f"- Deskripsi: {_text(item.get('description'))}",
                f"- Alasan: {_text(item.get('reason'))}",
                f"- Saran tindakan: {_text(item.get('suggested_action'))}",
                "",
            ]
        )
    return lines


def _slide_consistency_section(items: object) -> list[str]:
    values = [item for item in _as_list(items) if isinstance(item, dict)]
    lines = ["## Konsistensi Slide vs Laporan", ""]
    if not values:
        return lines + ["Belum ada data konsistensi slide.", ""]
    for item in values:
        lines.extend(
            [
                f"### Slide {item.get('slide_number', '-')}: {_text(item.get('slide_title'))}",
                "",
                f"- Klaim terdeteksi: {_text(item.get('detected_claim'))}",
                f"- Bagian laporan: {_text(item.get('matched_thesis_section'))}",
                f"- Status: {_text(item.get('status'))}",
                f"- Ringkasan: {_text(item.get('issue_summary'))}",
                f"- Saran perbaikan: {_text(item.get('suggested_fix'))}",
                f"- Evidence: {_text(item.get('evidence_excerpt'))}",
                "",
            ]
        )
    return lines


def _problematic_claims_section(items: object) -> list[str]:
    values = [item for item in _as_list(items) if isinstance(item, dict)]
    lines = ["## Klaim Bermasalah", ""]
    if not values:
        return lines + ["Tidak ada klaim bermasalah yang tersimpan.", ""]
    for index, item in enumerate(values, start=1):
        lines.extend(
            [
                f"### {index}. Slide {item.get('slide_number', '-')}",
                "",
                f"- Klaim: {_text(item.get('claim_text'))}",
                f"- Jenis masalah: {_text(item.get('problem_type'))}",
                f"- Risiko: {_text(item.get('risk_level'))}",
                f"- Evidence: {_text(item.get('evidence_excerpt'))}",
                f"- Saran revisi: {_text(item.get('suggested_revision'))}",
                "",
            ]
        )
    return lines


def _defense_questions_section(items: object) -> list[str]:
    values = [item for item in _as_list(items) if isinstance(item, dict)]
    lines = ["## Pertanyaan Penguji", ""]
    if not values:
        return lines + ["Belum ada pertanyaan penguji.", ""]
    for index, item in enumerate(values, start=1):
        lines.extend(
            [
                f"### {index}. {_text(item.get('category'))}",
                "",
                f"- Pertanyaan: {_text(item.get('question'))}",
                f"- Mengapa mungkin ditanya: {_text(item.get('why_asked'))}",
                f"- Panduan jawaban: {_text(item.get('answer_guidance'))}",
                f"- Bagian terkait: {_text(item.get('related_section'))}",
                f"- Tingkat kesulitan: {_text(item.get('difficulty'))}",
                "",
            ]
        )
    return lines


def _presentation_script_section(items: object) -> list[str]:
    values = [item for item in _as_list(items) if isinstance(item, dict)]
    lines = ["## Script Presentasi", ""]
    if not values:
        return lines + ["Belum ada script presentasi.", ""]
    for item in values:
        lines.extend(
            [
                f"### Slide {item.get('slide_number', '-')}: {_text(item.get('slide_title'))}",
                "",
                f"- Estimasi durasi: {item.get('estimated_duration_seconds', 0)} detik",
                "",
                _text(item.get("script_text")),
                "",
                "**Poin penting:**",
                "",
            ]
        )
        lines.extend([f"- {_text(point)}" for point in _as_list(item.get("key_points"))])
        lines.extend(["", "**Delivery tips:**", ""])
        lines.extend([f"- {_text(tip)}" for tip in _as_list(item.get("delivery_tips"))])
        lines.append("")
    return lines


def _related_text(item: dict[str, Any]) -> str:
    related = [item.get("related_chapter")]
    if item.get("related_slide"):
        related.append(f"Slide {item.get('related_slide')}")
    text = ", ".join(_text(value) for value in related if value)
    return text or "Tidak disebutkan"


def _as_list(value: object) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _text(value: object) -> str:
    if value is None:
        return "Tidak ditemukan dukungan eksplisit pada laporan yang diunggah."
    text = str(value).strip()
    return text or "Tidak ditemukan dukungan eksplisit pada laporan yang diunggah."


def _markdown_anchor(title: str) -> str:
    return (
        title.lower()
        .replace("&", "")
        .replace("/", "")
        .replace(" ", "-")
    )


def _bookmark_name(title: str) -> str:
    return f"toc_{_markdown_anchor(title).replace('-', '_')}"


def _strip_markdown_link(value: str) -> str:
    parsed = _parse_markdown_link(value)
    if parsed:
        return parsed[0]
    return value


def _parse_markdown_link(value: str) -> tuple[str, str] | None:
    if value.startswith("[") and "](" in value and value.endswith(")"):
        label, target = value.split("](", 1)
        return label.removeprefix("["), target.removesuffix(")")
    return None
